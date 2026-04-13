import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, size, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_formatted_text(paragraph, text, size=12):
    # This function parses simple markdown like **bold** and handles it as runs
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_formatted_text(p, text, 16 if level == 1 else 14)
    return p

def add_paragraph(doc, text, size=12, bold=False, align=None, indent=False, bullet=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    
    if align:
        p.alignment = align
    if indent and not bullet:
        p.paragraph_format.first_line_indent = Inches(0.5)
    
    if bullet:
        # Simple bullet handling
        p.style = 'List Bullet'
        text = text.lstrip('• ').lstrip('* ')

    add_formatted_text(p, text, size)
    return p

def add_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            tc = table.cell(i, j)
            p = tc.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            # Handle formatting in cells too
            add_formatted_text(p, str(cell), 10)
    return table

def generate_report():
    md_path = r"c:\Users\USER\Documents\AMBER-Project\Internship_Report_Content.md"
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    sections_md = content.split("---")

    for section_text in sections_md:
        lines = section_text.strip().split("\n")
        if not lines: continue
        
        doc.add_page_break()
        in_table = False
        table_data = []

        for line in lines:
            line = line.strip()
            if not line: continue

            if line.startswith("|"):
                if "---" in line: continue
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if cells:
                    table_data.append(cells)
                    in_table = True
                continue
            elif in_table:
                if table_data:
                    add_table(doc, table_data)
                table_data = []
                in_table = False

            if line.startswith("# "):
                add_paragraph(doc, line[2:].upper(), 16, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif line.startswith("## "):
                add_heading(doc, line[3:], 1)
            elif line.startswith("### "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                add_formatted_text(p, line[4:], 13)
                p.runs[0].bold = True # Ensure heading is bold
            elif line.startswith("* ") or line.startswith("• "):
                add_paragraph(doc, line, bullet=True)
            elif line.startswith("**") and line.endswith("**"):
                add_paragraph(doc, line.replace("**", ""), 12, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                add_paragraph(doc, line, 12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True)

        if in_table and table_data:
            add_table(doc, table_data)

    output_path = r"c:\Users\USER\Documents\AMBER-Project\AMBER-India_Final_40Page_Report_V2.docx"
    doc.save(output_path)
    print(f"Success! Final Clean Report generated at: {output_path}")

try:
    generate_report()
except Exception as e:
    print(f"Generation Error: {e}")
